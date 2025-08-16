"""
Google Drive Integration Service
Handles document generation and storage in Google Drive
"""

import os
import json
from typing import Optional, Dict, Any, List
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from io import BytesIO
from docx import Document
from docx.shared import Inches
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

class GoogleDriveService:
    def __init__(self):
        self.scopes = [
            'https://www.googleapis.com/auth/drive',
            'https://www.googleapis.com/auth/documents'
        ]
        self.service = None
        self.credentials = None
        self.client_secrets_file = '/app/backend/google_credentials.json'
        self.token_file = '/app/backend/token.json'
        
    def initialize_credentials(self) -> bool:
        """Initialize Google Drive credentials"""
        try:
            # Load existing token if available
            if os.path.exists(self.token_file):
                self.credentials = Credentials.from_authorized_user_file(
                    self.token_file, self.scopes
                )
            
            # If there are no valid credentials, return False for manual setup
            if not self.credentials or not self.credentials.valid:
                if self.credentials and self.credentials.expired and self.credentials.refresh_token:
                    try:
                        self.credentials.refresh(Request())
                    except Exception as e:
                        logger.error(f"Error refreshing credentials: {e}")
                        return False
                else:
                    # Need manual authorization
                    logger.warning("Google Drive credentials need to be set up")
                    return False
            
            # Save credentials for next run
            with open(self.token_file, 'w') as token:
                token.write(self.credentials.to_json())
                
            self.service = build('drive', 'v3', credentials=self.credentials)
            return True
            
        except Exception as e:
            logger.error(f"Error initializing Google Drive credentials: {e}")
            return False
    
    def get_authorization_url(self) -> str:
        """Get authorization URL for OAuth setup"""
        try:
            if not os.path.exists(self.client_secrets_file):
                raise FileNotFoundError("Google credentials file not found")
                
            flow = Flow.from_client_secrets_file(
                self.client_secrets_file,
                scopes=self.scopes,
                redirect_uri='http://localhost:8080/callback'
            )
            
            auth_url, _ = flow.authorization_url(
                access_type='offline',
                include_granted_scopes='true'
            )
            
            return auth_url
            
        except Exception as e:
            logger.error(f"Error getting authorization URL: {e}")
            return ""
    
    def exchange_code_for_token(self, code: str) -> bool:
        """Exchange authorization code for access token"""
        try:
            flow = Flow.from_client_secrets_file(
                self.client_secrets_file,
                scopes=self.scopes,
                redirect_uri='http://localhost:8080/callback'
            )
            
            flow.fetch_token(code=code)
            self.credentials = flow.credentials
            
            # Save credentials
            with open(self.token_file, 'w') as token:
                token.write(self.credentials.to_json())
                
            self.service = build('drive', 'v3', credentials=self.credentials)
            return True
            
        except Exception as e:
            logger.error(f"Error exchanging code for token: {e}")
            return False
    
    def create_client_folder(self, client_name: str) -> Optional[str]:
        """Create a folder for the client in Google Drive"""
        try:
            if not self.service:
                if not self.initialize_credentials():
                    raise Exception("Google Drive service not initialized")
            
            # Clean client name for folder
            folder_name = f"Cliente - {client_name}"
            
            # Check if folder already exists
            existing_folder = self.find_folder(folder_name)
            if existing_folder:
                return existing_folder['id']
            
            # Create new folder
            folder_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder'
            }
            
            folder = self.service.files().create(
                body=folder_metadata,
                fields='id'
            ).execute()
            
            logger.info(f"Created client folder: {folder_name} (ID: {folder.get('id')})")
            return folder.get('id')
            
        except Exception as e:
            logger.error(f"Error creating client folder: {e}")
            return None
    
    def find_folder(self, folder_name: str) -> Optional[Dict]:
        """Find a folder by name in Google Drive"""
        try:
            if not self.service:
                return None
                
            query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder'"
            results = self.service.files().list(
                q=query,
                fields="files(id, name)"
            ).execute()
            
            items = results.get('files', [])
            return items[0] if items else None
            
        except Exception as e:
            logger.error(f"Error finding folder: {e}")
            return None
    
    def get_template_document(self, template_name: str = "Template Procuração") -> Optional[str]:
        """Get template document from Google Drive"""
        try:
            if not self.service:
                if not self.initialize_credentials():
                    raise Exception("Google Drive service not initialized")
            
            # Search for template document
            query = f"name contains '{template_name}'"
            results = self.service.files().list(
                q=query,
                fields="files(id, name, mimeType)"
            ).execute()
            
            items = results.get('files', [])
            if not items:
                logger.warning(f"Template '{template_name}' not found")
                return None
            
            # Use first matching template
            template = items[0]
            
            # Download template content
            request = self.service.files().get_media(fileId=template['id'])
            file_content = BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            file_content.seek(0)
            return file_content.read().decode('utf-8')
            
        except Exception as e:
            logger.error(f"Error getting template document: {e}")
            return None
    
    def generate_procuracao(self, client_data: Dict[str, Any], process_data: Dict[str, Any] = None) -> Optional[Document]:
        """Generate power of attorney document"""
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('PROCURAÇÃO AD JUDICIA ET EXTRA JUDICIA', 0)
            title.alignment = 1  # Center alignment
            
            doc.add_paragraph("")
            
            # Client information
            client_info = f"""
OUTORGANTE: {client_data.get('name', '')}, {client_data.get('nationality', 'brasileiro(a)')}, 
{client_data.get('civil_status', '')}, {client_data.get('profession', '')}, 
portador do CPF nº {client_data.get('cpf', '')}, residente e domiciliado à 
{client_data.get('street', '')} nº {client_data.get('number', '')}, 
{client_data.get('district', '')}, {client_data.get('city', '')} - {client_data.get('state', '')}, 
CEP: {client_data.get('complement', '')}.
"""
            
            doc.add_paragraph(client_info.strip())
            doc.add_paragraph("")
            
            # Lawyer information (to be filled)
            lawyer_info = """
OUTORGADO: Dr(a). _________________, advogado(a), inscrito(a) na OAB/__ sob o nº _______, 
com escritório na _________________, CEP: _______, telefone: _______, 
e-mail: _________________.
"""
            
            doc.add_paragraph(lawyer_info.strip())
            doc.add_paragraph("")
            
            # Powers granted
            powers = """
Pelo presente instrumento, o(a) OUTORGANTE nomeia e constitui seu bastante procurador o(a) 
OUTORGADO acima qualificado, a quem confere os mais amplos poderes para o foro em geral, 
podendo propor contra quem de direito as ações que julgar convenientes, bem como defender 
o constituinte daquelas que lhe forem movidas, seguindo umas e outras até final decisão, 
usando dos recursos legais e acompanhando-os, conferindo-lhe ainda poderes especiais para:

• Transigir, desistir, firmar compromissos, fazer acordos judiciais e extrajudiciais;
• Receber e dar quitação;
• Substabelecer esta procuração, no todo ou em parte, com ou sem reserva de poderes;
• Requerer certidões, traslados e mais papéis;
• Assinar petições, contratos, escrituras e demais documentos;
• Representar o constituinte em todos os atos necessários ao cumprimento deste mandato.
"""
            
            doc.add_paragraph(powers.strip())
            doc.add_paragraph("")
            
            # Process information if provided
            if process_data:
                process_info = f"""
Esta procuração destina-se especificamente ao processo nº {process_data.get('process_number', '')}, 
do tipo {process_data.get('type', '')}, em tramitação no {process_data.get('court', 'tribunal competente')}.
"""
                doc.add_paragraph(process_info.strip())
                doc.add_paragraph("")
            
            # Date and signature
            current_date = datetime.now().strftime('%d de %B de %Y')
            signature_section = f"""
{client_data.get('city', 'Cidade')}, {current_date}.




_________________________________
{client_data.get('name', 'OUTORGANTE')}
CPF: {client_data.get('cpf', '')}




_________________________________
Dr(a). _________________
OAB/__ nº _______
OUTORGADO
"""
            
            doc.add_paragraph(signature_section.strip())
            
            return doc
            
        except Exception as e:
            logger.error(f"Error generating procuração: {e}")
            return None
    
    def save_document_to_drive(self, document: Document, filename: str, folder_id: str) -> Optional[str]:
        """Save document to Google Drive folder"""
        try:
            if not self.service:
                if not self.initialize_credentials():
                    raise Exception("Google Drive service not initialized")
            
            # Save document to temporary file
            temp_file = f"/tmp/{filename}"
            document.save(temp_file)
            
            # Upload to Google Drive
            file_metadata = {
                'name': filename,
                'parents': [folder_id]
            }
            
            media = MediaFileUpload(
                temp_file,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            file = self.service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()
            
            # Clean up temporary file
            os.remove(temp_file)
            
            logger.info(f"Document uploaded to Google Drive: {filename} (ID: {file.get('id')})")
            return file.get('webViewLink')
            
        except Exception as e:
            logger.error(f"Error saving document to Drive: {e}")
            return None
    
    def generate_and_save_procuracao(self, client_data: Dict[str, Any], process_data: Dict[str, Any] = None) -> Optional[str]:
        """Generate power of attorney and save to client's Google Drive folder"""
        try:
            # Create client folder
            client_name = client_data.get('name', 'Cliente')
            folder_id = self.create_client_folder(client_name)
            
            if not folder_id:
                raise Exception("Failed to create client folder")
            
            # Generate document
            document = self.generate_procuracao(client_data, process_data)
            if not document:
                raise Exception("Failed to generate document")
            
            # Create filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            process_suffix = f"_Processo_{process_data.get('process_number', 'GERAL')}" if process_data else "_GERAL"
            filename = f"Procuracao_{client_name.replace(' ', '_')}{process_suffix}_{timestamp}.docx"
            
            # Save to Google Drive
            drive_link = self.save_document_to_drive(document, filename, folder_id)
            
            return drive_link
            
        except Exception as e:
            logger.error(f"Error generating and saving procuração: {e}")
            return None
    
    def list_client_documents(self, client_name: str) -> List[Dict[str, Any]]:
        """List all documents in a client's folder"""
        try:
            if not self.service:
                if not self.initialize_credentials():
                    return []
            
            folder_name = f"Cliente - {client_name}"
            folder = self.find_folder(folder_name)
            
            if not folder:
                return []
            
            # List files in folder
            query = f"'{folder['id']}' in parents"
            results = self.service.files().list(
                q=query,
                fields="files(id, name, createdTime, webViewLink, mimeType)",
                orderBy="createdTime desc"
            ).execute()
            
            return results.get('files', [])
            
        except Exception as e:
            logger.error(f"Error listing client documents: {e}")
            return []
    
    def is_configured(self) -> bool:
        """Check if Google Drive integration is properly configured"""
        return (
            os.path.exists(self.client_secrets_file) and 
            self.initialize_credentials()
        )

# Global instance
google_drive_service = GoogleDriveService()